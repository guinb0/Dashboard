import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import numpy as np
from datetime import datetime
import sqlite3
import hashlib
import json
import os
from io import BytesIO

# É necessário instalar a biblioteca python-docx para gerar o relatório Word
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.shared import OxmlElement, qn
except ImportError:
    st.warning("⚠️ A biblioteca 'python-docx' não está instalada. A função de gerar relatórios em .docx estará desabilitada. Para habilitá-la, execute `pip install python-docx`.")
    Document = None

# Configuração da página
st.set_page_config(
    page_title="Dashboard de Avaliação de Riscos",
    page_icon="⚠️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Escalas de avaliação baseadas na metodologia SAROI
ESCALAS_IMPACTO = {
    "Muito baixo": {
        "valor": 1,
        "descricao": "Degradação de operações causando impactos mínimos nos objetivos"
    },
    "Baixo": {
        "valor": 2,
        "descricao": "Degradação de operações causando impactos pequenos nos objetivos"
    },
    "Médio": {
        "valor": 5,
        "descricao": "Interrupção de operações causando impactos significativos mas recuperáveis"
    },
    "Alto": {
        "valor": 8,
        "descricao": "Interrupção de operações causando impactos de reversão muito difícil"
    },
    "Muito alto": {
        "valor": 10,
        "descricao": "Paralisação de operações causando impactos irreversíveis/catastróficos"
    }
}

ESCALAS_PROBABILIDADE = {
    "Muito baixa": {
        "valor": 1,
        "descricao": "Evento improvável de ocorrer. Não há elementos que indiquem essa possibilidade"
    },
    "Baixa": {
        "valor": 2,
        "descricao": "Evento raro de ocorrer. Poucos elementos indicam essa possibilidade"
    },
    "Média": {
        "valor": 5,
        "descricao": "Evento possível de ocorrer. Elementos indicam moderadamente essa possibilidade"
    },
    "Alta": {
        "valor": 8,
        "descricao": "Evento provável de ocorrer. Elementos indicam consistently essa possibilidade"
    },
    "Muito alta": {
        "valor": 10,
        "descricao": "Evento praticamente certo de ocorrer. Elementos indicam claramente essa possibilidade"
    }
}

# Modalidades de mitigação padrão (baseadas na planilha fornecida)
MODALIDADES_PADRAO = [
    "Permuta por imóvel já construído",
    "Permuta por edificação a construir (terreno terceiros)",
    "Permuta por obra (terreno da União)",
    "Build to Suit (terreno da União)",
    "Contratação com dação em pagamento",
    "Obra pública convencional"
]

# Aspectos a serem considerados para cada risco (extraídos da planilha) - AMPLIADO
ASPECTOS_RISCOS = {
    'Descumprimento do Prazo de entrega': {
        'impacto': [
            "Condições de segurança/conservação do imóvel utilizado pelo órgão",
            "Custo de locação do imóvel utilizado pelo órgão", 
            "Taxa de ocupação do imóvel utilizado pelo órgão",
            "Impacto na continuidade dos serviços públicos",
            "Custos adicionais com prorrogações contratuais"
        ],
        'probabilidade': [
            "Estrutura de monitoramento e mecanismos contratuais de sanção previstos",
            "Complexidade técnica do empreendimento e riscos externos (licenças, clima, logística)",
            "Grau de maturidade dos projetos disponibilizados",
            "Características do local de implantação",
            "Histórico de cumprimento de prazos da empresa contratada",
            "Capacidade técnica e financeira do contratado"
        ]
    },
    'Indisponibilidade de imóveis públicos p/ implantação ou dação em permuta': {
        'impacto': [
            "Quantidade de imóveis disponíveis e nível de desembraço desses imóveis",
            "Impacto na viabilidade econômica da operação",
            "Necessidade de recursos orçamentários adicionais",
            "Comprometimento da estratégia de otimização do patrimônio público"
        ],
        'probabilidade': [
            "Quantidade de imóveis disponíveis e nível de desembraço desses imóveis",
            "Processos judiciais em andamento sobre os imóveis",
            "Situação registral e documental dos imóveis",
            "Interesse de outros órgãos públicos nos mesmo imóveis",
            "Complexidade dos procedimentos de desafetação"
        ]
    },
    'Condições de mercado desfavoráveis': {
        'impacto': [
            "Condições de segurança/conservação do imóvel utilizado pelo órgão",
            "Custo de locação do imóvel utilizado pelo órgão",
            "Taxa de ocupação do imóvel utilizado pelo órgão",
            "Redução da competitividade no processo licitatório",
            "Aumento dos custos da operação"
        ],
        'probabilidade': [
            "Valor do investimento necessário (valor imóveis x torna x construção)",
            "Atratividade dos lotes ofertados (valor; possibilidades de utilização; tendências do mercado)",
            "Grau de especialização exigida do investidor",
            "Grau de aquecimento do mercado x taxa de juros x rentabilidade esperada",
            "Manifestações de interesse ou consultas públicas realizadas",
            "Histórico de certames semelhantes e nível de participação",
            "Cenário econômico nacional e setorial"
        ]
    },
    'Abandono da obra pela empresa': {
        'impacto': [
            "Condições de segurança/conservação do imóvel utilizado pelo órgão",
            "Custo de locação do imóvel utilizado pelo órgão",
            "Taxa de ocupação do imóvel utilizado pelo órgão",
            "Custos de nova licitação e retomada da obra",
            "Atraso significativo na entrega do empreendimento"
        ],
        'probabilidade': [
            "Requisitos técnicos e financeiros a serem previstos no processo de seleção",
            "Garantias contratuais e outras salvaguardas a serem previstas",
            "Garantias contratuais e outras salvaguardas previstas na modelagem",
            "Percentual do novo prédio a ser ocupado pela Administração",
            "Situação financeira e histórico da empresa contratada",
            "Robustez dos mecanismos de acompanhamento da execução"
        ]
    },
    'Baixa rentabilização do estoque de imóveis': {
        'impacto': [
            "Valor dos imóveis dados em permuta na operação frente ao valor do imóvel adquirido",
            "Amplitude do potencial de valorização dos imóveis dados em permuta",
            "Grau de contribuição da operação para a redução de imóveis ociosos",
            "Prejuízo patrimonial para a União",
            "Redução da eficiência da gestão patrimonial"
        ],
        'probabilidade': [
            "Grau de contribuição da operação para a redução de imóveis ociosos",
            "Adequação do uso proposto às características do imóvel",
            "Potencial de economia de despesas (Eficiência do plano de gestão do ativo)",
            "Eficiência do plano de gestão do ativo pós-permuta",
            "Demanda do mercado e probabilidade de valorização dos imóveis",
            "Localização e características dos imóveis ofertados",
            "Estratégia de alienação ou exploração econômica"
        ]
    },
    'Dotação orçamentária insuficiente': {
        'impacto': [
            "Condições de segurança/conservação do imóvel utilizado pelo órgão",
            "Custo de locação do imóvel utilizado pelo órgão",
            "Taxa de ocupação do imóvel utilizado pelo órgão",
            "Percentual do valor da operação que será custeada com recursos orçamentários",
            "Inviabilização completa do projeto",
            "Necessidade de renegociação contratual"
        ],
        'probabilidade': [
            "Peso da previsão de despesa em relação à dotação orçamentária de investimento",
            "Informações constantes da LOA e PPA",
            "Histórico de contingenciamento do órgão",
            "Peso político dos órgãos beneficiários",
            "Cenário fiscal e orçamentário da União",
            "Priorização do projeto no planejamento governamental"
        ]
    },
    'Questionamento jurídico': {
        'impacto': [
            "Paralisação completa ou parcial do projeto",
            "Custos adicionais com defesa jurídica",
            "Perda de credibilidade institucional",
            "Necessidade de reformulação da modelagem",
            "Impacto na continuidade dos serviços públicos"
        ],
        'probabilidade': [
            "Complexidade e inovação da modelagem jurídica adotada",
            "Precedentes jurisprudenciais sobre modalidades similares",
            "Robustez da fundamentação legal da contratação",
            "Histórico de questionamentos em projetos similares",
            "Atuação de órgãos de controle externo",
            "Transparência e aderência aos princípios da administração pública",
            "Qualidade da documentação jurídica do processo"
        ]
    },
    'Baixa qualidade dos serviços entregues': {
        'impacto': [
            "Custos adicionais com reparos e adequações",
            "Insatisfação dos usuários finais",
            "Redução da vida útil do empreendimento",
            "Necessidade de nova contratação para correções",
            "Comprometimento da imagem institucional",
            "Impacto na funcionalidade operacional"
        ],
        'probabilidade': [
            "Rigor dos critérios de qualificação técnica",
            "Estrutura de fiscalização e acompanhamento técnico",
            "Especificações técnicas e padrões de qualidade definidos",
            "Histórico de qualidade dos serviços da empresa contratada",
            "Mecanismos contratuais de garantia de qualidade",
            "Complexidade técnica dos serviços demandados",
            "Adequação entre o preço contratado e o padrão de qualidade esperado"
        ]
    }
}

# Funções para gerenciamento do banco de dados
def init_db():
    """Inicializa o banco de dados SQLite"""
    conn = sqlite3.connect('riscos.db')
    c = conn.cursor()
    
    # Tabela de usuários
    c.execute('''CREATE TABLE IF NOT EXISTS usuarios
                 (id INTEGER PRIMARY KEY AUTOINCREMENT,
                  username TEXT UNIQUE NOT NULL,
                  password_hash TEXT NOT NULL)''')
    
    # Tabela de logs de ações
    c.execute('''CREATE TABLE IF NOT EXISTS logs
                 (id INTEGER PRIMARY KEY AUTOINCREMENT,
                  timestamp DATETIME DEFAULT CURRENT_TIMESTAMP,
                  username TEXT NOT NULL,
                  acao TEXT NOT NULL,
                  detalhes TEXT)''')
    
    # Inserir usuários padrão se não existirem
    usuarios_padrao = [
        ("SPU 1", hashlib.sha256("1234".encode()).hexdigest()),
        ("SPU 2", hashlib.sha256("1234".encode()).hexdigest()),
        ("SPU 3", hashlib.sha256("1234".encode()).hexdigest())
    ]
    
    for usuario, senha_hash in usuarios_padrao:
        try:
            c.execute("INSERT INTO usuarios (username, password_hash) VALUES (?, ?)", 
                      (usuario, senha_hash))
        except sqlite3.IntegrityError:
            pass  # Usuário já existe
    
    conn.commit()
    conn.close()

def verificar_login(username, password):
    """Verifica se as credenciais são válidas"""
    conn = sqlite3.connect('riscos.db')
    c = conn.cursor()
    
    password_hash = hashlib.sha256(password.encode()).hexdigest()
    c.execute("SELECT * FROM usuarios WHERE username = ? AND password_hash = ?", 
              (username, password_hash))
    
    resultado = c.fetchone()
    conn.close()
    
    return resultado is not None

def registrar_acao(username, acao, detalhes=None):
    """Registra uma ação no log"""
    conn = sqlite3.connect('riscos.db')
    c = conn.cursor()
    
    c.execute("INSERT INTO logs (username, acao, detalhes) VALUES (?, ?, ?)",
              (username, acao, json.dumps(detalhes) if detalhes else None))
    
    conn.commit()
    conn.close()

def obter_logs():
    """Obtém todos os logs do sistema"""
    conn = sqlite3.connect('riscos.db')
    c = conn.cursor()
    
    c.execute("SELECT timestamp, username, acao, detalhes FROM logs ORDER BY timestamp DESC")
    logs = c.fetchall()
    conn.close()
    
    return logs

def classificar_risco(valor_risco):
    """Classifica o risco baseado no valor calculado"""
    if valor_risco <= 10:
        return "Baixo", "#28a745"
    elif valor_risco <= 25:
        return "Médio", "#ffc107"
    else:
        return "Alto", "#dc3545"

def gerar_relatorio_word():
    """Gera relatório completo e amplo em formato Word"""
    if Document is None:
        st.error("📋 A biblioteca python-docx não está instalada. Não é possível gerar o relatório.")
        return None
        
    try:
        # Obter nome do projeto da session_state
        nome_projeto = st.session_state.get('nome_projeto', 'Projeto')
        
        # Criar documento
        doc = Document()
        
        # Título principal com nome do projeto
        title = doc.add_heading(f'RELATÓRIO EXECUTIVO DE AVALIAÇÃO DE RISCOS - {nome_projeto}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtítulo
        subtitle = doc.add_heading("Metodologia - Análise Comparativa de Modalidades de Contratação", level=1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # NOVO: Nome do Projeto como título dentro do documento
        doc.add_heading(f"Projeto: {nome_projeto}", level=2)
        doc.add_paragraph()
        
        # NOVA SEÇÃO: Informações do responsável pelo relatório
        # Solicitar identificação do usuário se não estiver definida
        if 'identificacao_relatorio' not in st.session_state or st.session_state.identificacao_relatorio is None:
            st.session_state.identificacao_relatorio = {
                'nome': st.session_state.user,
                'unidade': 'Unidade Padrão',
                'orgao': 'SPU',
                'email': 'usuario@spu.gov.br'
            }
        
        # Informações do relatório com identificação
        info_para = doc.add_paragraph()
        info_para.add_run("Data da Análise: ").bold = True
        info_para.add_run(f"{datetime.now().strftime('%d/%m/%Y às %H:%M')}")
        info_para.add_run("\nMetodologia: ").bold = True
        info_para.add_run("Roteiro de Auditoria de Gestão de Riscos - SAROI")
        info_para.add_run("\nVersão do Sistema: ").bold = True
        info_para.add_run("2.0 - Análise Ampliada")
        
        # Adicionar informações do responsável
        info_para.add_run("\n\nRESPONSÁVEL PELA ANÁLISE:").bold = True
        info_para.add_run(f"\nNome: {st.session_state.identificacao_relatorio['nome']}")
        info_para.add_run(f"\nDivisão: {st.session_state.identificacao_relatorio['unidade']}")
        if st.session_state.identificacao_relatorio['orgao']:
            info_para.add_run(f"\nÓrgão: {st.session_state.identificacao_relatorio['orgao']}")
        if st.session_state.identificacao_relatorio['email']:
            info_para.add_run(f"\nE-mail: {st.session_state.identificacao_relatorio['email']}")
        
        doc.add_paragraph()
        
        # 1. RESUMO EXECUTIVO
        doc.add_heading('1. RESUMO EXECUTIVO', level=1)
        
        total_riscos = len(st.session_state.riscos)
        riscos_altos = sum(1 for r in st.session_state.riscos if r['classificacao'] == 'Alto')
        riscos_medios = sum(1 for r in st.session_state.riscos if r['classificacao'] == 'Médio')
        riscos_baixos = sum(1 for r in st.session_state.riscos if r['classificacao'] == 'Baixo')
        risco_inerente_total = sum(r['risco_inerente'] for r in st.session_state.riscos)
        
        # Calcular melhor e pior modalidade
        risco_acumulado_por_modalidade = {}
        for modalidade in st.session_state.modalidades:
            risco_residual_total = 0
            for risco in st.session_state.riscos:
                if modalidade in risco['modalidades']:
                    fator_mitigacao = risco['modalidades'][modalidade]
                    risco_residual = risco['risco_inerente'] * fator_mitigacao
                    risco_residual_total += risco_residual
            risco_acumulado_por_modalidade[modalidade] = risco_residual_total
        
        melhor_modalidade = min(risco_acumulado_por_modalidade.keys(), 
                               key=lambda x: risco_acumulado_por_modalidade[x])
        pior_modalidade = max(risco_acumulado_por_modalidade.keys(), 
                             key=lambda x: risco_acumulado_por_modalidade[x])
        
        resumo = f"""
        Este relatório apresenta análise quantitativa de {total_riscos} riscos identificados para o projeto 
        . A análise inclui avaliação detalhada 
        de impacto e probabilidade, cálculo de riscos inerentes e residuais, e comparação sistemática entre 
        {len(st.session_state.modalidades)} modalidades de contratação.
        
        MÉTRICAS PRINCIPAIS DO PROJETO:
        • Total de Riscos Analisados: {total_riscos}
        • Distribuição por Classificação:
          - Riscos ALTOS: {riscos_altos} ({riscos_altos/total_riscos*100:.1f}%)
          - Riscos MÉDIOS: {riscos_medios} ({riscos_medios/total_riscos*100:.1f}%)
          - Riscos BAIXOS: {riscos_baixos} ({riscos_baixos/total_riscos*100:.1f}%)
        • Risco Inerente Total: {risco_inerente_total:.1f} pontos
        
        RESULTADO DA ANÁLISE COMPARATIVA:
        • MODALIDADE RECOMENDADA: {melhor_modalidade}
          - Risco Residual: {risco_acumulado_por_modalidade[melhor_modalidade]:.1f} pontos
        • MODALIDADE DE MAIOR RISCO: {pior_modalidade}
          - Risco Residual: {risco_acumulado_por_modalidade[pior_modalidade]:.1f} pontos
        • DIFERENÇA DE RISCO: {risco_acumulado_por_modalidade[pior_modalidade] - risco_acumulado_por_modalidade[melhor_modalidade]:.1f} pontos
        """
        doc.add_paragraph(resumo)
        
        # 2. METODOLOGIA DETALHADA
        doc.add_heading('2. METODOLOGIA E CRITÉRIOS DE AVALIAÇÃO', level=1)
        metodologia = """
        A avaliação seguiu rigorosamente a metodologia estabelecida no "Roteiro de Auditoria de 
        Gestão de Riscos", aplicando escalas quantitativas padronizadas e critérios objetivos.
        
        2.1 ESCALAS DE AVALIAÇÃO
        
        IMPACTO (Consequências para os objetivos):
        • Muito baixo (1): Degradação mínima das operações
        • Baixo (2): Degradação pequena, facilmente recuperável
        • Médio (5): Interrupção significativa mas recuperável
        • Alto (8): Interrupção grave, reversão muito difícil
        • Muito alto (10): Paralisação com impactos irreversíveis
        
        PROBABILIDADE (Chance de ocorrência):
        • Muito baixa (1): Evento improvável, sem elementos indicativos
        • Baixa (2): Evento raro, poucos elementos indicam possibilidade
        • Média (5): Evento possível, elementos moderadamente indicativos
        • Alta (8): Evento provável, elementos consistentemente indicativos
        • Muito alta (10): Evento praticamente certo, elementos claramente indicativos
        
        2.2 CÁLCULO DO RISCO INERENTE
        
        O risco inerente é calculado pela multiplicação: IMPACTO × PROBABILIDADE
        
        2.3 CLASSIFICAÇÃO DOS RISCOS
        
        • BAIXO: Risco inerente ≤ 10 pontos
        • MÉDIO: Risco inerente entre 11 e 25 pontos  
        • ALTO: Risco inerente > 25 pontos
        
        2.4 CÁLCULO DO RISCO RESIDUAL
        
        Para cada modalidade, o risco residual é calculado aplicando-se o fator de mitigação:
        RISCO RESIDUAL = RISCO INERENTE × FATOR DE MITIGAÇÃO
        
        Onde o fator de mitigação varia de 0,0 (elimina totalmente o risco) a 1,0 (não mitiga o risco).
        """
        doc.add_paragraph(metodologia)
        
        # 3. ANÁLISE DETALHADA DOS RISCOS
        doc.add_heading('3. ANÁLISE DETALHADA DOS RISCOS IDENTIFICADOS', level=1)
        
        for i, risco in enumerate(st.session_state.riscos, 1):
            doc.add_heading(f'3.{i} {risco["risco_chave"]}', level=2)
            
            # Avaliação quantitativa
            aval_para = doc.add_paragraph()
            aval_para.add_run("\nAVALIAÇÃO QUANTITATIVA:").bold = True
            aval_para.add_run(f"\n• Impacto: {risco['impacto_valor']} ({risco['impacto_nivel']})\n")
            aval_para.add_run("Justificativa do risco: ").bold = True
            aval_para.add_run(risco.get("descricao", ""))
            aval_para.add_run(f"\n\n• Probabilidade: {risco['probabilidade_valor']} ({risco['probabilidade_nivel']})\n")
            aval_para.add_run("Justificativa de Probabilidade de ocorrência: ").bold = True
            aval_para.add_run(risco.get("contexto_especifico", ""))
            aval_para.add_run(f"\n\n• Risco Inerente: {risco['risco_inerente']} pontos")
            aval_para.add_run(f"\n• Classificação: {risco['classificacao']}")
            
            # Análise por modalidade - AGORA EM TABELA
            doc.add_heading('3.x Análise por Modalidade', level=3)
            
            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'
            
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Modalidade'
            hdr_cells[1].text = 'Fator de Mitigação'
            hdr_cells[2].text = 'Risco Residual'
            hdr_cells[3].text = 'Eficácia (%)'
            hdr_cells[4].text = 'Justificativa'
            
            justificativas_modalidades = risco.get("justificativas_modalidades", {})
            for modalidade, fator in risco['modalidades'].items():
                risco_residual = risco['risco_inerente'] * fator
                eficacia = (1 - fator) * 100
                
                row_cells = table.add_row().cells
                row_cells[0].text = modalidade
                row_cells[1].text = f"{fator:.1f}"
                row_cells[2].text = f"{risco_residual:.1f} ({classificar_risco(risco_residual)[0]})"
                row_cells[3].text = f"{eficacia:.1f}%"
                row_cells[4].text = justificativas_modalidades.get(modalidade, "")
        
        # 4. ANÁLISE COMPARATIVA DAS MODALIDADES
        doc.add_heading('4. ANÁLISE COMPARATIVA DAS MODALIDADES', level=1)
        
        # Calcular dados comparativos detalhados
        dados_comparativos = {}
        
        for modalidade in st.session_state.modalidades:
            risco_residual_total = 0
            risco_inerente_aplicavel = 0
            count_riscos = 0
            
            for risco in st.session_state.riscos:
                if modalidade in risco['modalidades']:
                    fator_mitigacao = risco['modalidades'][modalidade]
                    risco_residual = risco['risco_inerente'] * fator_mitigacao
                    risco_residual_total += risco_residual
                    risco_inerente_aplicavel += risco['risco_inerente']
                    count_riscos += 1
            
            eficacia_total = ((risco_inerente_aplicavel - risco_residual_total) / risco_inerente_aplicavel * 100) if risco_inerente_aplicavel > 0 else 0
            
            dados_comparativos[modalidade] = {
                'risco_residual_total': risco_residual_total,
                'risco_inerente_aplicavel': risco_inerente_aplicavel,
                'eficacia_percentual': eficacia_total,
                'classificacao': classificar_risco(risco_residual_total)[0],
                'riscos_aplicaveis': count_riscos
            }
        
        # Tabela comparativa principal
        doc.add_heading('4.1 Quadro Comparativo Consolidado', level=2)
        
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Ranking'
        hdr_cells[1].text = 'Modalidade'
        hdr_cells[2].text = 'Risco Residual Total'
        hdr_cells[3].text = 'Eficácia Mitigação (%)'
        hdr_cells[4].text = 'Classificação Final'
        hdr_cells[5].text = 'Riscos Aplicáveis'
        
        # Ordenar modalidades por risco residual
        modalidades_ordenadas = sorted(dados_comparativos.items(), 
                                       key=lambda x: x[1]['risco_residual_total'])
        
        for i, (modalidade, dados) in enumerate(modalidades_ordenadas, 1):
            row_cells = table.add_row().cells
            row_cells[0].text = f"{i}º"
            row_cells[1].text = modalidade
            row_cells[2].text = f"{dados['risco_residual_total']:.1f}"
            row_cells[3].text = f"{dados['eficacia_percentual']:.1f}%"
            row_cells[4].text = dados['classificacao']
            row_cells[5].text = f"{dados['riscos_aplicaveis']}/{total_riscos}"
        
        # 4.2 Análise de Performance
        doc.add_heading('4.2 Análise de Performance por Modalidade', level=2)
        
        for i, (modalidade, dados) in enumerate(modalidades_ordenadas, 1):
            posicao_texto = "RECOMENDADA" if i == 1 else "NÃO RECOMENDADA" if i == len(modalidades_ordenadas) else f"{i}ª COLOCADA"
            
            performance_para = doc.add_paragraph()
            performance_para.add_run(f"{modalidade} - {posicao_texto}").bold = True
            performance_para.add_run(f"""
            • Risco Residual Total: {dados['risco_residual_total']:.1f} pontos
            • Eficácia de Mitigação: {dados['eficacia_percentual']:.1f}%
            • Classificação de Risco: {dados['classificacao']}
            • Redução Absoluta do Risco: {dados['risco_inerente_aplicavel'] - dados['risco_residual_total']:.1f} pontos
            • Riscos Aplicáveis: {dados['riscos_aplicaveis']} de {total_riscos} riscos
            """)
        
        # 5. MATRIZ DETALHADA DE RISCOS
        doc.add_heading('5. MATRIZ DETALHADA DE RISCOS POR MODALIDADE', level=1)
        
        # Criar tabela expandida
        num_cols = 3 + len(st.session_state.modalidades)
        table = doc.add_table(rows=1, cols=num_cols)
        table.style = 'Table Grid'
        
        # Cabeçalhos
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Risco'
        hdr_cells[1].text = 'Impacto'
        hdr_cells[2].text = 'Probabilidade'
        for i, modalidade in enumerate(st.session_state.modalidades):
            col_name = modalidade[:15] + "..." if len(modalidade) > 15 else modalidade
            hdr_cells[3 + i].text = col_name
        
        # Dados por risco
        for risco in st.session_state.riscos:
            row_cells = table.add_row().cells
            risco_name = risco['risco_chave'][:25] + "..." if len(risco['risco_chave']) > 25 else risco['risco_chave']
            row_cells[0].text = risco_name
            row_cells[1].text = str(risco['impacto_valor'])
            row_cells[2].text = str(risco['probabilidade_valor'])
            
            for i, modalidade in enumerate(st.session_state.modalidades):
                if modalidade in risco['modalidades']:
                    fator = risco['modalidades'][modalidade]
                    risco_residual = risco['risco_inerente'] * fator
                    row_cells[3 + i].text = f"{risco_residual:.1f}"
                else:
                    row_cells[3 + i].text = "N/A"
        
        # Linha de totais
        row_cells = table.add_row().cells
        row_cells[0].text = "TOTAL ACUMULADO"
        row_cells[1].text = "-"
        row_cells[2].text = "-"
        
        for i, modalidade in enumerate(st.session_state.modalidades):
            if modalidade in dados_comparativos:
                row_cells[3 + i].text = f"{dados_comparativos[modalidade]['risco_residual_total']:.1f}"
            else:
                row_cells[3 + i].text = "N/A"
        
        # 6. RECOMENDAÇÕES E CONCLUSÕES
        doc.add_heading('6. RECOMENDAÇÕES EXECUTIVAS', level=1)
        
        melhor_modalidade_dados = dados_comparativos[melhor_modalidade]
        pior_modalidade_dados = dados_comparativos[pior_modalidade]
        
        recomendacoes = f"""
        6.1 MODALIDADE RECOMENDADA
        
        Com base na análise quantitativa realizada, recomenda-se a adoção da modalidade:
        "{melhor_modalidade}"
        
        JUSTIFICATIVAS TÉCNICAS:
        • Menor risco residual acumulado: {melhor_modalidade_dados['risco_residual_total']:.1f} pontos
        • Maior eficácia de mitigação: {melhor_modalidade_dados['eficacia_percentual']:.1f}%
        • Classificação de Risco: {melhor_modalidade_dados['classificacao']}
        • Redução Absoluta do Risco: {melhor_modalidade_dados['risco_inerente_aplicavel'] - melhor_modalidade_dados['risco_residual_total']:.1f} pontos
        • Riscos Aplicáveis: {melhor_modalidade_dados['riscos_aplicaveis']} de {total_riscos} riscos
        
        6.2 MODALIDADES NÃO RECOMENDADAS
        
        A modalidade de maior risco identificada é:
        "{pior_modalidade}"
        
        RAZÕES PARA NÃO RECOMENDAÇÃO:
        • Maior risco residual acumulado: {pior_modalidade_dados['risco_residual_total']:.1f} pontos
        • Menor eficácia de mitigação: {pior_modalidade_dados['eficacia_percentual']:.1f}%
        • Classificação de Risco: {pior_modalidade_dados['classificacao']}
        
        6.3 IMPACTO DA ESCOLHA DA MODALIDADE
        
        A diferença entre a melhor e pior modalidade é de {pior_modalidade_dados['risco_residual_total'] - melhor_modalidade_dados['risco_residual_total']:.1f} pontos de risco, 
        representando {(pior_modalidade_dados['risco_residual_total'] - melhor_modalidade_dados['risco_residual_total'])/risco_inerente_total*100:.1f}% 
        do risco total do projeto.
        
        Esta diferença demonstra a importância crítica da escolha adequada da modalidade de contratação 
        para o sucesso do empreendimento.
        """
        doc.add_paragraph(recomendacoes)
        
        # 7. CONCLUSÕES FINAIS
        doc.add_heading('7. CONCLUSÕES E CONSIDERAÇÕES FINAIS', level=1)
        
        conclusoes = f"""
        A presente análise, baseada na metodologia consolidada do SAROI, permitiu uma avaliação 
        objetiva e fundamentada das modalidades de contratação disponíveis para o projeto.
        
        PRINCIPAIS RESULTADOS:
        
        1. RISCO TOTAL DO PROJETO: {risco_inerente_total:.1f} pontos (antes da mitigação)
        
        2. ESTRATÉGIA ÓTIMA IDENTIFICADA: {melhor_modalidade}
            - Reduz o risco total para {melhor_modalidade_dados['risco_residual_total']:.1f} pontos
            - Eficácia de mitigação de {melhor_modalidade_dados['eficacia_percentual']:.1f}%
            - Redução absoluta de {melhor_modalidade_dados['risco_inerente_aplicavel'] - melhor_modalidade_dados['risco_residual_total']:.1f} pontos de risco
            
        3. AMPLITUDE DE VARIAÇÃO: As modalidades analisadas apresentam variação de risco residual 
            de {pior_modalidade_dados['risco_residual_total'] - melhor_modalidade_dados['risco_residual_total']:.1f} pontos, 
            evidenciando a relevância da escolha estratégica.
            
        4. CONFORMIDADE METODOLÓGICA: A análise seguiu integralmente os preceitos estabelecidos 
            pelo SAROI para gestão de riscos em projetos públicos, garantindo objetividade e 
            fundamentação técnica para a tomada de decisão.
        
        CONSIDERAÇÕES PARA IMPLEMENTAÇÃO:
        
        • A modalidade recomendada deve ser implementada observando-se os aspectos específicos 
          identificados na análise de cada risco.
        • Recomenda-se o monitoramento contínuo dos fatores de risco durante a execução do projeto.
        • Os resultados desta análise devem ser revisados caso ocorram mudanças significativas 
          no contexto do projeto ou nas condições de mercado.
        
        Esta análise fornece base técnica sólida e metodologicamente consistente para a tomada 
        de decisão, em total conformidade com as melhores práticas de gestão de riscos estabelecidas 
        pelos órgãos de controle.
        """
        doc.add_paragraph(conclusoes)
        
        # ANEXOS
        doc.add_heading('ANEXOS', level=1)
        
        # Anexo I - Escalas utilizadas
        doc.add_heading('ANEXO I - Escalas de Avaliação Utilizadas', level=2)
        
        escalas_texto = """
        ESCALA DE IMPACTO:
        1 - Muito baixo: Degradação de operações causando impactos mínimos nos objetivos
        2 - Baixo: Degradação de operações causando impactos pequenos nos objetivos  
        5 - Médio: Interrupção de operações causando impactos significativos mas recuperáveis
        8 - Alto: Interrupção de operações causando impactos de reversão muito difícil
        10 - Muito alto: Paralisação de operações causando impactos irreversíveis/catastróficos
        
        ESCALA DE PROBABILIDADE:
        1 - Muito baixa: Evento improvável de ocorrer. Não há elementos que indiquem essa possibilidade
        2 - Baixa: Evento raro de ocorrer. Poucos elementos indicam essa possibilidade
        5 - Média: Evento possível de ocorrer. Elementos indicam moderadamente essa possibilidade  
        8 - Alta: Evento provável de ocorrer. Elementos indicam consistentemente essa possibilidade
        10 - Muito alta: Evento praticamente certo de ocorrer. Elementos indicam claramente essa possibilidade
        """
        doc.add_paragraph(escalas_texto)
        
        # Rodapé
        doc.add_paragraph()
        doc.add_paragraph("_" * 50)
        rodape = doc.add_paragraph()
        rodape.add_run("Relatório gerado automaticamente pelo Sistema de Avaliação de Riscos SAROI v2.0").italic = True
        rodape.add_run(f"\nData e hora: {datetime.now().strftime('%d/%m/%Y às %H:%M')}")
        rodape.add_run(f"\nResponsável: {st.session_state.identificacao_relatorio['nome']} - {st.session_state.identificacao_relatorio['unidade']}")
        if st.session_state.identificacao_relatorio['orgao']:
            rodape.add_run(f"\nDivisão: {st.session_state.identificacao_relatorio['unidade']}")
        rodape.add_run(f"\nTotal de páginas estimadas: {len(doc.paragraphs) // 20 + 1}")
        
        # Salvar em buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer
        
    except ImportError:
        st.error("📋 Para gerar relatórios Word, instale a biblioteca python-docx: pip install python-docx")
        return None
    except Exception as e:
        st.error(f"Erro ao gerar relatório: {str(e)}")
        return None

def calcular_risco_inerente(impacto, probabilidade):
    """Calcula o risco inerente (Impacto x Probabilidade)"""
    return impacto * probabilidade

def criar_heatmap_modalidades_melhorado(riscos_comparacao):
    """Cria heatmap melhorado com mais clareza visual"""
    
    # Preparar dados para o heatmap
    dados_heatmap = []
    modalidades = st.session_state.modalidades
    
    for risco in riscos_comparacao:
        linha_risco = []
        for modalidade in modalidades:
            if modalidade in risco['modalidades']:
                fator_mitigacao = risco['modalidades'][modalidade]
                risco_residual = risco['risco_inerente'] * fator_mitigacao
                linha_risco.append(risco_residual)
            else:
                linha_risco.append(0)
        dados_heatmap.append(linha_risco)
    
    # Criar labels mais limpos
    labels_riscos = []
    for risco in riscos_comparacao:
        nome = risco['risco_chave']
        if len(nome) > 30:
            # Tentar quebrar em palavras-chave
            palavras = nome.split()
            if len(palavras) > 3:
                nome = " ".join(palavras[:3]) + "..."
            else:
                nome = nome[:30] + "..."
        labels_riscos.append(nome)
    
    labels_modalidades = []
    for mod in modalidades:
        if len(mod) > 25:
            # Abreviar modalidades longas
            if "Permuta" in mod:
                nome = mod.replace("Permuta por ", "P.").replace(" (terreno", "(t.")
            elif "Build to Suit" in mod:
                nome = "Build to Suit (União)"
            elif "Contratação" in mod:
                nome = "Contrat. c/ dação"
            else:
                nome = mod[:25] + "..."
        else:
            nome = mod
        labels_modalidades.append(nome)
    
    # Criar figura com customização melhorada
    fig = go.Figure(data=go.Heatmap(
        z=dados_heatmap,
        x=labels_modalidades,
        y=labels_riscos,
        colorscale=[
            [0.0, '#00ff00'],    # Verde para risco zero/muito baixo
            [0.3, '#90EE90'],    # Verde claro
            [0.5, '#ffff00'],    # Amarelo para risco médio
            [0.7, '#FFA500'],    # Laranja
            [1.0, '#ff0000']     # Vermelho para risco alto
        ],
        showscale=True,
        colorbar=dict(
            title="Risco Residual",
            tickmode="linear",
            tick0=0,
            dtick=10
        ),
        text=[[f"{val:.1f}" if val > 0 else "0" for val in linha] for linha in dados_heatmap],
        texttemplate="%{text}",
        textfont={"size": 10, "color": "black"},
        hoverongaps=False,
        hovertemplate="<b>%{y}</b><br>" +
                      "Modalidade: %{x}<br>" +
                      "Risco Residual: %{z:.1f}<br>" +
                      "<extra></extra>"
    ))
    
    # Melhorar layout
    fig.update_layout(
        title={
            'text': "Mapa de Calor: Risco Residual por Modalidade",
            'x': 0.5,
            'xanchor': 'center',
            'font': {'size': 16, 'color': 'darkblue'}
        },
        xaxis_title="Modalidades de Contratação",
        yaxis_title="Riscos Identificados",
        width=1000,
        height=700,
        font=dict(size=11),
        xaxis=dict(tickangle=45, side="bottom"),
        yaxis=dict(autorange="reversed"),  # Inverter ordem para melhor leitura
        margin=dict(l=200, r=100, t=100, b=150)
    )
    
    return fig

def criar_heatmap_eficacia_melhorado(riscos_comparacao):
    """Cria heatmap de eficácia melhorado"""
    
    # Preparar dados para eficácia
    dados_eficacia = []
    modalidades = st.session_state.modalidades
    
    for risco in riscos_comparacao:
        linha_eficacia = []
        for modalidade in modalidades:
            if modalidade in risco['modalidades']:
                fator_mitigacao = risco['modalidades'][modalidade]
                eficacia = (1 - fator_mitigacao) * 100  # Percentual de redução
                linha_eficacia.append(eficacia)
            else:
                linha_eficacia.append(0)
        dados_eficacia.append(linha_eficacia)
    
    # Usar os mesmos labels do heatmap anterior para consistência
    labels_riscos = []
    for risco in riscos_comparacao:
        nome = risco['risco_chave']
        if len(nome) > 30:
            palavras = nome.split()
            if len(palavras) > 3:
                nome = " ".join(palavras[:3]) + "..."
            else:
                nome = nome[:30] + "..."
        labels_riscos.append(nome)
    
    labels_modalidades = []
    for mod in modalidades:
        if len(mod) > 25:
            if "Permuta" in mod:
                nome = mod.replace("Permuta por ", "P.").replace(" (terreno", "(t.")
            elif "Build to Suit" in mod:
                nome = "Build to Suit (União)"
            elif "Contratação" in mod:
                nome = "Contrat. c/ dação"
            else:
                nome = mod[:25] + "..."
        else:
            nome = mod
        labels_modalidades.append(nome)
    
    # Criar figura
    fig = go.Figure(data=go.Heatmap(
        z=dados_eficacia,
        x=labels_modalidades,
        y=labels_riscos,
        colorscale='RdYlGn',  # Vermelho-Amarelo-Verde (invertido para eficácia)
        showscale=True,
        colorbar=dict(
            title="Eficácia (%)",
            tickmode="linear",
            tick0=0,
            dtick=20
        ),
        text=[[f"{val:.0f}%" if val > 0 else "0%" for val in linha] for linha in dados_eficacia],
        texttemplate="%{text}",
        textfont={"size": 10, "color": "black"},
        hoverongaps=False,
        hovertemplate="<b>%{y}</b><br>" +
                      "Modalidade: %{x}<br>" +
                      "Eficácia: %{z:.1f}%<br>" +
                      "<extra></extra>"
    ))
    
    # Layout
    fig.update_layout(
        title={
            'text': "Mapa de Calor: Eficácia de Mitigação por Modalidade",
            'x': 0.5,
            'xanchor': 'center',
            'font': {'size': 16, 'color': 'darkblue'}
        },
        xaxis_title="Modalidades de Contratação",
        yaxis_title="Riscos Identificados",
        width=1000,
        height=700,
        font=dict(size=11),
        xaxis=dict(tickangle=45, side="bottom"),
        yaxis=dict(autorange="reversed"),
        margin=dict(l=200, r=100, t=100, b=150)
    )
    
    return fig

def inicializar_dados():
    """Inicializa os dados padrão do sistema"""
    if 'riscos' not in st.session_state:
        riscos_iniciais = [
            {
                'risco_chave': 'Descumprimento do Prazo de entrega',
                'descricao': 'Risco de a empresa contratada não cumprir o prazo de entrega da obra ou serviço, gerando atrasos e possíveis prejuízos para a Administração Pública.',
                'impacto_nivel': 'Alto',
                'impacto_valor': 8,
                'probabilidade_nivel': 'Alta',
                'probabilidade_valor': 8,
                'risco_inerente': 64,
                'classificacao': 'Alto',
                'modalidades': {
                    'Permuta por imóvel já construído': 0.6,
                    'Permuta por edificação a construir (terreno terceiros)': 0.6,
                    'Permuta por obra (terreno da União)': 0.4,
                    'Build to Suit (terreno da União)': 0.4,
                    'Contratação com dação em pagamento': 0.8,
                    'Obra pública convencional': 0.2
                }
            },
            {
                'risco_chave': 'Indisponibilidade de imóveis públicos p/ implantação ou dação em permuta',
                'descricao': 'Risco de não haver imóveis públicos disponíveis ou adequados para a implantação de projetos ou para serem utilizados como dação em pagamento em operações de permuta.',
                'impacto_nivel': 'Médio',
                'impacto_valor': 5,
                'probabilidade_nivel': 'Média',
                'probabilidade_valor': 5,
                'risco_inerente': 25,
                'classificacao': 'Médio',
                'modalidades': {
                    'Permuta por imóvel já construído': 0.2,
                    'Permuta por edificação a construir (terreno terceiros)': 0.2,
                    'Permuta por obra (terreno da União)': 0.6,
                    'Build to Suit (terreno da União)': 0.6,
                    'Contratação com dação em pagamento': 0.4,
                    'Obra pública convencional': 1.0
                }
            },
            {
                'risco_chave': 'Condições de mercado desfavoráveis',
                'descricao': 'Risco de as condições de mercado (ex: taxas de juros elevadas, baixa demanda) inviabilizarem ou encarecerem a operação de contratação ou permuta.',
                'impacto_nivel': 'Médio',
                'impacto_valor': 5,
                'probabilidade_nivel': 'Média',
                'probabilidade_valor': 5,
                'risco_inerente': 25,
                'classificacao': 'Médio',
                'modalidades': {
                    'Permuta por imóvel já construído': 0.4,
                    'Permuta por edificação a construir (terreno terceiros)': 0.4,
                    'Permuta por obra (terreno da União)': 0.6,
                    'Build to Suit (terreno da União)': 0.6,
                    'Contratação com dação em pagamento': 0.2,
                    'Obra pública convencional': 0.8
                }
            },
            {
                'risco_chave': 'Abandono da obra pela empresa',
                'descricao': 'Risco de a empresa contratada abandonar a obra ou serviço antes da conclusão, gerando a necessidade de nova licitação e atrasos significativos.',
                'impacto_nivel': 'Alto',
                'impacto_valor': 8,
                'probabilidade_nivel': 'Baixa',
                'probabilidade_valor': 2,
                'risco_inerente': 16,
                'classificacao': 'Médio',
                'modalidades': {
                    'Permuta por imóvel já construído': 0.8,
                    'Permuta por edificação a construir (terreno terceiros)': 0.8,
                    'Permuta por obra (terreno da União)': 0.4,
                    'Build to Suit (terreno da União)': 0.4,
                    'Contratação com dação em pagamento': 0.2,
                    'Obra pública convencional': 0.2
                }
            },
            {
                'risco_chave': 'Baixa rentabilização do estoque de imóveis',
                'descricao': 'Impacto total, somente superável no caso de a SPU disponibilizar diversos imóveis de alto interesse pelo mercado.',
                'impacto_nivel': 'Alto',
                'impacto_valor': 8,
                'probabilidade_nivel': 'Alta',
                'probabilidade_valor': 8,
                'risco_inerente': 64,
                'classificacao': 'Alto',
                'modalidades': {
                    'Permuta por imóvel já construído': 1.0,
                    'Permuta por edificação a construir (terreno terceiros)': 1.0,
                    'Permuta por obra (terreno da União)': 0.2,
                    'Build to Suit (terreno da União)': 0.6,
                    'Contratação com dação em pagamento': 0.4,
                    'Obra pública convencional': 0.8
                }
            },
            {
                'risco_chave': 'Dotação orçamentária insuficiente',
                'descricao': 'Impacto total, somente superável no caso de a SPU disponibilizar diversos imóveis de alto interesse pelo mercado.',
                'impacto_nivel': 'Muito alto',
                'impacto_valor': 10,
                'probabilidade_nivel': 'Muito alta',
                'probabilidade_valor': 10,
                'risco_inerente': 100,
                'classificacao': 'Alto',
                'modalidades': {
                    'Permuta por imóvel já construído': 0.0,
                    'Permuta por edificação a construir (terreno terceiros)': 0.1,
                    'Permuta por obra (terreno da União)': 0.1,
                    'Build to Suit (terreno da União)': 0.4,
                    'Contratação com dação em pagamento': 0.4,
                    'Obra pública convencional': 1.0
                }
            },
            {
                'risco_chave': 'Questionamento jurídico',
                'descricao': 'Possibilidade de questionamentos jurídicos quanto à legalidade da modalidade de contratação escolhida, especialmente em modalidades inovadoras ou complexas.',
                'impacto_nivel': 'Médio',
                'impacto_valor': 5,
                'probabilidade_nivel': 'Média',
                'probabilidade_valor': 5,
                'risco_inerente': 25,
                'classificacao': 'Médio',
                'modalidades': {
                    'Permuta por imóvel já construído': 0.2,
                    'Permuta por edificação a construir (terreno terceiros)': 0.4,
                    'Permuta por obra (terreno da União)': 0.4,
                    'Build to Suit (terreno da União)': 0.4,
                    'Contratação com dação em pagamento': 0.6,
                    'Obra pública convencional': 0.1
                }
            },
            {
                'risco_chave': 'Baixa qualidade dos serviços entregues',
                'descricao': 'Risco de que os serviços ou obras entregues não atendam aos padrões de qualidade exigidos, comprometendo a funcionalidade e durabilidade do empreendimento.',
                'impacto_nivel': 'Médio',
                'impacto_valor': 5,
                'probabilidade_nivel': 'Baixa',
                'probabilidade_valor': 2,
                'risco_inerente': 10,
                'classificacao': 'Médio',
                'modalidades': {
                    'Permuta por imóvel já construído': 0.8,
                    'Permuta por edificação a construir (terreno terceiros)': 0.8,
                    'Permuta por obra (terreno da União)': 0.4,
                    'Build to Suit (terreno da União)': 0.4,
                    'Contratação com dação em pagamento': 0.2,
                    'Obra pública convencional': 0.2
                }
            }
        ]
        
        # Garante que a chave 'justificativas_modalidades' e 'contexto_especifico' exista em todos os riscos
        textos_exemplo_prob = [
            "A probabilidade é alta devido à complexidade da obra e do terreno.",
            "A probabilidade é média, pois o histórico de projetos similares na região é misto.",
            "A probabilidade é baixa, já que o contrato prevê mecanismos de fiscalização rigorosos.",
            "A probabilidade é muito alta, dado o cenário econômico atual e as manifestações de interesse já recebidas."
        ]
        
        textos_exemplo_mitigacao = [
            "A mitigação por esta modalidade se deve à menor dependência de terceiros.",
            "Esta modalidade é eficaz porque permite maior controle sobre a qualidade dos materiais.",
            "O fator de mitigação é baixo devido à alta volatilidade do mercado para este tipo de ativo.",
            "A escolha desta modalidade reduz o risco de abandono da obra, pois o pagamento está atrelado à entrega.",
            "Justificativa para o fator: a modalidade transfere a maior parte da responsabilidade para o parceiro privado.",
            "O risco residual é alto nesta modalidade, pois a Administração assume os custos de renegociação.",
            "Esta é a modalidade mais segura em termos de orçamento, pois o custo é fixo e previamente estabelecido.",
            "O fator de mitigação reflete o controle limitado da Administração sobre a execução da obra neste modelo."
        ]
        
        for risco in riscos_iniciais:
            if "justificativas_modalidades" not in risco:
                risco["justificativas_modalidades"] = {
                    modalidade: np.random.choice(textos_exemplo_mitigacao) for modalidade in risco["modalidades"]
                }
            if "contexto_especifico" not in risco or not risco["contexto_especifico"]:
                risco["contexto_especifico"] = np.random.choice(textos_exemplo_prob)
        
        st.session_state.riscos = riscos_iniciais
        
    if 'modalidades' not in st.session_state:
        st.session_state.modalidades = MODALIDADES_PADRAO.copy()

def cadastro_riscos():
    st.header("📝 Cadastro de Riscos")
    
    if st.session_state.riscos:
        st.info(f"💡 **{len(st.session_state.riscos)} riscos** da planilha já estão carregados. Use o formulário abaixo para adicionar novos riscos.")
        st.info("💡 **Dica:** Para personalizar riscos existentes conforme seu caso concreto, use a aba '✏️ Editar Riscos'")
    
    with st.form("cadastro_risco"):
        col1, col2 = st.columns(2)
        
        with col1:
            risco_chave = st.text_input(
                "Risco-Chave:",
                placeholder="Ex: Descumprimento do Prazo de entrega"
            )
            
            descricao_risco = st.text_area(
                "Descrição/Justificativa do Risco:",
                placeholder="Descreva os aspectos que levam a este risco..."
            )
            
            contexto_especifico = st.text_area(
                "Justificativa de mudança de Probabilidade:",
                placeholder="Ex: Localização, tipo de obra, prazo, complexidade...",
                help="Aspectos específicos do seu projeto que influenciam este risco"
            )
        
        with col2:
            # Avaliação de Impacto
            st.subheader("🎯 Avaliação de Impacto")
            impacto_nivel = st.selectbox(
                "Nível de Impacto:",
                list(ESCALAS_IMPACTO.keys()),
                help="Selecione o nível de impacto baseado na escala SAROI"
            )
            st.info(f"**{impacto_nivel}** (Valor: {ESCALAS_IMPACTO[impacto_nivel]['valor']})")
            st.caption(ESCALAS_IMPACTO[impacto_nivel]['descricao'])
            
            # Avaliação de Probabilidade
            st.subheader("📊 Avaliação de Probabilidade")
            probabilidade_nivel = st.selectbox(
                "Nível de Probabilidade:",
                list(ESCALAS_PROBABILIDADE.keys()),
                help="Selecione o nível de probabilidade baseado na escala SAROI"
            )
            st.info(f"**{probabilidade_nivel}** (Valor: {ESCALAS_PROBABILIDADE[probabilidade_nivel]['valor']})")
            st.caption(ESCALAS_PROBABILIDADE[probabilidade_nivel]['descricao'])
        
        # Cálculo automático do risco inerente
        impacto_valor = ESCALAS_IMPACTO[impacto_nivel]['valor']
        probabilidade_valor = ESCALAS_PROBABILIDADE[probabilidade_nivel]['valor']
        risco_inerente = calcular_risco_inerente(impacto_valor, probabilidade_valor)
        classificacao, cor = classificar_risco(risco_inerente)
        
        st.divider()
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Impacto", impacto_valor)
        with col2:
            st.metric("Probabilidade", probabilidade_valor)
        with col3:
            st.metric("Risco Inerente", f"{risco_inerente} - {classificacao}")
        
        # Avaliação por modalidade
        st.subheader("🔄 Avaliação por Modalidades")
        st.info("Para cada modalidade, defina os fatores de mitigação (0.0 = elimina totalmente o risco, 1.0 = não mitiga)")
        
        modalidades_avaliacao = {}
        justificativas_modalidades = {}
        cols = st.columns(min(3, len(st.session_state.modalidades)))
        
        for i, modalidade in enumerate(st.session_state.modalidades):
            with cols[i % len(cols)]:
                fator = st.slider(
                    f"{modalidade}:",
                    min_value=0.0,
                    max_value=1.0,
                    value=0.0,  # ALTERADO: Valor padrão agora é 0.0
                    step=0.1,
                    key=f"modalidade_{i}"
                )
                justificativa = st.text_area("Justificativa:", key=f"justificativa_{i}")
                modalidades_avaliacao[modalidade] = fator
                justificativas_modalidades[modalidade] = justificativa
                
                # Calcular risco residual
                risco_residual = risco_inerente * fator
                class_residual, _ = classificar_risco(risco_residual)
                st.caption(f"Risco Residual: {risco_residual:.1f} ({class_residual})")
        
        submitted = st.form_submit_button("💾 Salvar Risco", type="primary")
        
        if submitted and risco_chave:
            novo_risco = {
                'risco_chave': risco_chave,
                'descricao': descricao_risco,
                'contexto_especifico': contexto_especifico,
                'impacto_nivel': impacto_nivel,
                'impacto_valor': impacto_valor,
                'probabilidade_nivel': probabilidade_nivel,
                'probabilidade_valor': probabilidade_valor,
                'risco_inerente': risco_inerente,
                'classificacao': classificacao,
                'modalidades': modalidades_avaliacao,
                'justificativas_modalidades': justificativas_modalidades,
                'personalizado': True,
                'criado_por': st.session_state.user,
                'data_criacao': datetime.now().strftime("%d/%m/%Y %H:%M")
            }
            
            st.session_state.riscos.append(novo_risco)
            
            # Registrar a ação no log
            registrar_acao(
                st.session_state.user, 
                "Criou risco", 
                {"risco": risco_chave, "detalhes": novo_risco}
            )
            
            st.success(f"✅ Risco '{risco_chave}' salvo com sucesso!")
            st.rerun()

def editar_riscos():
    st.header("✏️ Editar Riscos Existentes")
    
    if not st.session_state.riscos:
        st.warning("⚠️ Nenhum risco cadastrado para editar. Vá para a aba 'Cadastro de Riscos' para adicionar riscos.")
        return
    
    st.info("💡 **Personalize a avaliação** dos riscos conforme as características específicas do seu caso concreto.")
    
    # Seleção do risco para editar
    col1, col2 = st.columns([2, 1])
    
    with col1:
        opcoes_riscos = [f"{i+1}. {risco['risco_chave']}" for i, risco in enumerate(st.session_state.riscos)]
        risco_selecionado_str = st.selectbox(
            "Selecione o risco para editar:",
            opcoes_riscos,
            help="Escolha o risco que deseja personalizar"
        )
    
    with col2:
        if st.button("🔄 Recarregar página"):
            st.rerun()
    
    if not risco_selecionado_str:
        return
    
    # Extrair índice do risco selecionado
    indice_risco = int(risco_selecionado_str.split('.')[0]) - 1
    risco_atual = st.session_state.riscos[indice_risco]
    
    # Mostrar informações atuais do risco
    with st.expander("📋 Informações atuais do risco", expanded=True):
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Impacto Atual", f"{risco_atual['impacto_valor']} ({risco_atual['impacto_nivel']})")
        with col2:
            st.metric("Probabilidade Atual", f"{risco_atual['probabilidade_valor']} ({risco_atual['probabilidade_nivel']})")
        with col3:
            st.metric("Risco Inerente Atual", f"{risco_atual['risco_inerente']} ({risco_atual['classificacao']})")
    
    # Formulário de edição
    with st.form(f"editar_risco_{indice_risco}"):
        col1, col2 = st.columns(2)

        with col1:
            st.subheader("🎯 Reavaliação de Impacto")
            st.caption("Considere as características específicas do seu projeto:")
            risco_nome = risco_atual['risco_chave']
            if risco_nome in ASPECTOS_RISCOS:
                with st.expander("💡 Aspectos a serem considerados para IMPACTO", expanded=True):
                    st.write("**Considere os seguintes aspectos ao avaliar o impacto:**")
                    for i, aspecto in enumerate(ASPECTOS_RISCOS[risco_nome]['impacto'], 1):
                        st.write(f"• {aspecto}")
                    st.info("💡 **Dica:** Analise como cada aspecto se aplica ao seu caso específico antes de definir o nível de impacto.")
            niveis_impacto = list(ESCALAS_IMPACTO.keys())
            indice_impacto_atual = niveis_impacto.index(risco_atual['impacto_nivel'])
            novo_impacto_nivel = st.selectbox(
                "Novo Nível de Impacto:",
                niveis_impacto,
                index=indice_impacto_atual,
                help="Baseado nas características do seu caso concreto"
            )
            with st.expander("📖 Consultar Escala de Impacto"):
                for nivel, dados in ESCALAS_IMPACTO.items():
                    emoji = "👉" if nivel == novo_impacto_nivel else "•"
                    st.write(f"{emoji} **{nivel}** (Valor: {dados['valor']}): {dados['descricao']}")

        with col2:
            st.subheader("📊 Reavaliação de Probabilidade")
            st.caption("Considere a realidade do seu contexto:")
            if risco_nome in ASPECTOS_RISCOS:
                with st.expander("💡 Aspectos a serem considerados para PROBABILIDADE", expanded=True):
                    st.write("**Considere os seguintes aspectos ao avaliar a probabilidade:**")
                    for i, aspecto in enumerate(ASPECTOS_RISCOS[risco_nome]['probabilidade'], 1):
                        st.write(f"• {aspecto}")
                    st.info("💡 **Dica:** Analise como cada aspecto se aplica ao seu contexto antes de definir o nível de probabilidade.")
            niveis_probabilidade = list(ESCALAS_PROBABILIDADE.keys())
            indice_probabilidade_atual = niveis_probabilidade.index(risco_atual['probabilidade_nivel'])
            nova_probabilidade_nivel = st.selectbox(
                "Novo Nível de Probabilidade:",
                niveis_probabilidade,
                index=indice_probabilidade_atual,
                help="Baseado na realidade do seu contexto"
            )
            with st.expander("📖 Consultar Escala de Probabilidade"):
                for nivel, dados in ESCALAS_PROBABILIDADE.items():
                    emoji = "👉" if nivel == nova_probabilidade_nivel else "•"
                    st.write(f"{emoji} **{nivel}** (Valor: {dados['valor']}): {dados['descricao']}")

        st.subheader("Justificativas")
        col1, col2 = st.columns(2)

        with col1:
            nova_descricao = st.text_area(
                "Justificativa fator de impacto:",
                value=risco_atual.get('descricao', ''),
                help="Descreva as características específicas do seu caso que justificam a avaliação"
            )

        with col2:
            contexto_especifico = st.text_area(
                "Fatores específicos que influenciam a probabilidade deste risco:",
                value=risco_atual.get('contexto_especifico', ''),
                placeholder="Ex: Localização, tipo de obra, prazo, complexidade, recursos disponíveis...",
                help="Descreva os aspectos únicos do seu projeto"
            )
        
        # Calcular novos valores
        novo_impacto_valor = ESCALAS_IMPACTO[novo_impacto_nivel]['valor']
        nova_probabilidade_valor = ESCALAS_PROBABILIDADE[nova_probabilidade_nivel]['valor']
        novo_risco_inerente = calcular_risco_inerente(novo_impacto_valor, nova_probabilidade_valor)
        nova_classificacao, _ = classificar_risco(novo_risco_inerente)
        
        # Mostrar comparação
        st.divider()
        st.subheader("📊 Comparação: Antes vs Depois")
        
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric(
                "Impacto", 
                f"{novo_impacto_valor} ({novo_impacto_nivel})",
                delta=f"{novo_impacto_valor - risco_atual['impacto_valor']}"
            )
        with col2:
            st.metric(
                "Probabilidade", 
                f"{nova_probabilidade_valor} ({nova_probabilidade_nivel})",
                delta=f"{nova_probabilidade_valor - risco_atual['probabilidade_valor']}"
            )
        with col3:
            st.metric(
                "Risco Inerente", 
                f"{novo_risco_inerente} ({nova_classificacao})",
                delta=f"{novo_risco_inerente - risco_atual['risco_inerente']}"
            )
        
        # Reavaliação das modalidades
        st.subheader("🔄 Reavaliação das Modalidades")
        st.info("Ajuste os fatores de mitigação considerando as características específicas do seu caso")
        
        novas_modalidades = {}
        novas_justificativas = {}
        cols = st.columns(min(3, len(st.session_state.modalidades)))
        
        for i, modalidade in enumerate(st.session_state.modalidades):
            with cols[i % len(cols)]:
                valor_atual = risco_atual['modalidades'].get(modalidade, 0.5)
                novo_fator = st.slider(
                    f"{modalidade}:",
                    min_value=0.0,
                    max_value=1.0,
                    value=valor_atual,
                    step=0.1,
                    key=f"editar_modalidade_{indice_risco}_{i}"
                )
                
                # Acessa com .get() para evitar KeyError
                justificativa_modalidade = risco_atual.get("justificativas_modalidades", {}).get(modalidade, "")
                nova_justificativa = st.text_area(
                    "Justificativa:",
                    value=justificativa_modalidade,
                    key=f"justificativa_modalidade_{indice_risco}_{i}"
                )
                novas_modalidades[modalidade] = novo_fator
                novas_justificativas[modalidade] = nova_justificativa
                
                # Mostrar comparação do risco residual
                risco_residual_antigo = risco_atual['risco_inerente'] * valor_atual
                risco_residual_novo = novo_risco_inerente * novo_fator
                delta_residual = risco_residual_novo - risco_residual_antigo
                
                st.caption(f"Risco Residual: {risco_residual_novo:.1f}")
                if delta_residual != 0:
                    st.caption(f"Δ: {delta_residual:+.1f}")
        
        submitted = st.form_submit_button("💾 Salvar Alterações", type="primary")
        
        if submitted:
            # Atualizar o risco
            st.session_state.riscos[indice_risco].update({
                'impacto_nivel': novo_impacto_nivel,
                'impacto_valor': novo_impacto_valor,
                'probabilidade_nivel': nova_probabilidade_nivel,
                'probabilidade_valor': nova_probabilidade_valor,
                'risco_inerente': novo_risco_inerente,
                'classificacao': nova_classificacao,
                'descricao': nova_descricao,
                'contexto_especifico': contexto_especifico,
                'modalidades': novas_modalidades,
                'justificativas_modalidades': novas_justificativas,
                'editado': True,
                'data_edicao': datetime.now().strftime("%d/%m/%Y %H:%M")
            })
            
            # Registrar a ação no log
            registrar_acao(
                st.session_state.user, 
                "Editou risco", 
                {"risco": risco_atual['risco_chave'], "alteracoes": {
                    "impacto_anterior": risco_atual['impacto_valor'],
                    "impacto_novo": novo_impacto_valor,
                    "probabilidade_anterior": risco_atual['probabilidade_valor'],
                    "probabilidade_novo": nova_probabilidade_valor,
                    "risco_inerente_anterior": risco_atual['risco_inerente'],
                    "risco_inerente_novo": novo_risco_inerente
                }}
            )
            
            st.success(f"✅ Risco '{risco_atual['risco_chave']}' atualizado com sucesso!")
            st.rerun()

def analise_riscos():
    st.header("📊 Análise de Riscos")
    
    if not st.session_state.riscos:
        st.warning("⚠️ Nenhum risco cadastrado para análise.")
        return
    
    # Filtros
    st.subheader("🔍 Filtros")
    col1, col2, col3 = st.columns(3)
    
    with col1:
        filtro_classificacao = st.multiselect(
            "Filtrar por Classificação:",
            ["Alto", "Médio", "Baixo"],
            default=["Alto", "Médio", "Baixo"]
        )
    
    with col2:
        filtro_busca = st.text_input(
            "Buscar por nome do risco:",
            placeholder="Digite parte do nome..."
        )
    
    with col3:
        filtro_tipo = st.selectbox(
            "Filtrar por Tipo:",
            ["Todos", "Originais da planilha", "Personalizados", "Adicionados"]
        )
    
    # Aplicar filtros
    riscos_filtrados = []
    for risco in st.session_state.riscos:
        # Filtro por classificação
        if risco['classificacao'] not in filtro_classificacao:
            continue
        
        # Filtro por busca
        if filtro_busca and filtro_busca.lower() not in risco['risco_chave'].lower():
            continue
        
        # Filtro por tipo
        if filtro_tipo == "Originais da planilha" and (risco.get('personalizado', False) or risco.get('editado', False)):
            continue
        elif filtro_tipo == "Personalizados" and not risco.get('editado', False):
            continue
        elif filtro_tipo == "Adicionados" and not risco.get('personalizado', False):
            continue
        
        riscos_filtrados.append(risco)
    
    if not riscos_filtrados:
        st.warning("Nenhum risco encontrado com os filtros aplicados.")
        return
    
    # Visualizações
    col1, col2 = st.columns(2)
    
    with col1:
        # Gráfico de distribuição por classificação
        classificacoes = [r['classificacao'] for r in riscos_filtrados]
        df_class = pd.DataFrame({'Classificação': classificacoes})
        contagem_class = df_class['Classificação'].value_counts()
        
        fig_pizza = px.pie(
            values=contagem_class.values,
            names=contagem_class.index,
            title="Distribuição de Riscos por Classificação",
            color_discrete_map={"Baixo": "#28a745", "Médio": "#ffc107", "Alto": "#dc3545"}
        )
        st.plotly_chart(fig_pizza, use_container_width=True)
    
    with col2:
        # Gráfico de dispersão Impacto x Probabilidade
        df_scatter = pd.DataFrame(riscos_filtrados)
        
        # Adicionar indicador de personalização
        df_scatter['Tipo'] = df_scatter.apply(
            lambda row: 'Personalizado' if row.get('editado', False)
            else ('Adicionado' if row.get('personalizado', False) else 'Original'), 
            axis=1
        )
        
        fig_scatter = px.scatter(
            df_scatter,
            x='probabilidade_valor',
            y='impacto_valor',
            size='risco_inerente',
            color='Tipo',
            hover_data=['risco_chave', 'classificacao'],
            title="Matriz de Riscos (Impacto x Probabilidade)",
            labels={'probabilidade_valor': 'Probabilidade', 'impacto_valor': 'Impacto'},
            color_discrete_map={
                "Original": "#6c757d", 
                "Personalizado": "#007bff", 
                "Adicionado": "#28a745"
            }
        )
        fig_scatter.update_layout(xaxis_range=[0, 11], yaxis_range=[0, 11])
        st.plotly_chart(fig_scatter, use_container_width=True)
    
    # Tabela detalhada
    st.subheader("📋 Detalhamento dos Riscos")
    
    # Preparar dados para a tabela
    dados_tabela = []
    for risco in riscos_filtrados:
        # Indicador de tipo
        if risco.get('editado', False):
            tipo_icon = "✏️"
            tipo = "Personalizado"
        elif risco.get('personalizado', False):
            tipo_icon = "➕"
            tipo = "Adicionado"
        else:
            tipo_icon = "📋"
            tipo = "Original"
        
        dados_tabela.append({
            'Tipo': f"{tipo_icon} {tipo}",
            'Risco': risco['risco_chave'],
            'Impacto': risco['impacto_valor'],
            'Probabilidade': risco['probabilidade_valor'],
            'Risco Inerente': risco['risco_inerente'],
            'Classificação': risco['classificacao']
        })
    
    df_display = pd.DataFrame(dados_tabela)
    
    # Exibir tabela
    try:
        st.dataframe(df_display, use_container_width=True)
    except:
        st.dataframe(df_display, use_container_width=True)
    
    # Legenda
    st.caption("**Legenda:** 📋 Original da planilha | ✏️ Personalizado para caso concreto | ➕ Adicionado manualmente")
    
    # Detalhes expandidos para riscos personalizados
    riscos_editados_detalhes = [r for r in riscos_filtrados if r.get('editado', False)]
    if riscos_editados_detalhes:
        with st.expander(f"🔍 Detalhes dos {len(riscos_editados_detalhes)} riscos personalizados"):
            for risco in riscos_editados_detalhes:
                st.write(f"**{risco['risco_chave']}**")
                if 'contexto_especifico' in risco and risco['contexto_especifico']:
                    st.write(f"*Contexto específico:* {risco['contexto_especifico']}")
                if 'data_edicao' in risco:
                    st.write(f"*Última edição:* {risco['data_edicao']}")
                if 'descricao' in risco and risco['descricao']:
                    st.write(f"*Descrição:* {risco['descricao']}")
                st.write("---")
    
    # Nova seção: Visão rápida do risco residual acumulado
    st.subheader("⚡ Visão Rápida - Risco Residual por Modalidade")
    st.caption("Baseado nos riscos filtrados atualmente")
    
    # Calcular risco residual para os riscos filtrados
    if len(riscos_filtrados) > 1:  # Só mostrar se há mais de um risco
        risco_residual_rapido = {}
        
        for modalidade in st.session_state.modalidades:
            risco_total = 0
            count = 0
            
            for risco in riscos_filtrados:
                if modalidade in risco['modalidades']:
                    risco_residual = risco['risco_inerente'] * risco['modalidades'][modalidade]
                    risco_total += risco_residual
                    count += 1
            
            if count > 0:
                risco_residual_rapido[modalidade] = risco_total
        
        if risco_residual_rapido:
            # Mostrar as 3 melhores e 3 piores modalidades
            modalidades_ordenadas = sorted(risco_residual_rapido.items(), key=lambda x: x[1])
            
            col1, col2 = st.columns(2)
            
            with col1:
                st.success("**🏆 Melhores Modalidades (Menor Risco Residual):**")
                for i, (modalidade, risco) in enumerate(modalidades_ordenadas[:3], 1):
                    classificacao, _ = classificar_risco(risco)
                    st.write(f"{i}. **{modalidade}**: {risco:.1f} ({classificacao})")
            
            with col2:
                st.error("**⚠️ Modalidades de Maior Risco Residual:**")
                for i, (modalidade, risco) in enumerate(reversed(modalidades_ordenadas[-3:]), 1):
                    classificacao, _ = classificar_risco(risco)
                    st.write(f"{i}. **{modalidade}**: {risco:.1f} ({classificacao})")
            
            st.info(f"💡 **Dica:** Para análise completa do risco acumulado, acesse a aba '🔄 Comparação de Modalidades'")
    else:
        st.info("💡 Selecione múltiplos riscos para ver a análise de risco residual acumulado.")

def comparacao_modalidades():
    st.header("🔄 Comparação de Modalidades")
    
    if not st.session_state.riscos:
        st.warning("⚠️ Nenhum risco cadastrado para comparação.")
        return
    
    # Selecionar riscos para comparação
    riscos_opcoes = [f"{i+1}. {r['risco_chave']}" for i, r in enumerate(st.session_state.riscos)]
    riscos_selecionados = st.multiselect(
        "Selecione os riscos para comparação:",
        riscos_opcoes,
        default=riscos_opcoes
    )
    
    if not riscos_selecionados:
        st.warning("Selecione pelo menos um risco para comparação.")
        return
    
    # Extrair índices dos riscos selecionados
    indices_selecionados = [int(r.split('.')[0]) - 1 for r in riscos_selecionados]
    riscos_comparacao = [st.session_state.riscos[i] for i in indices_selecionados]
    
    # Calcular risco residual ACUMULADO por modalidade
    st.subheader("📊 Risco Residual Acumulado por Modalidade")
    st.info("💡 **Risco Residual Acumulado** = Soma de todos os riscos residuais para cada modalidade. Representa o risco total ao escolher uma estratégia.")
    
    risco_acumulado_por_modalidade = {}
    risco_inerente_total = sum(risco['risco_inerente'] for risco in riscos_comparacao)
    
    # Calcular dados detalhados e acumulados
    dados_comparacao = []
    
    for modalidade in st.session_state.modalidades:
        risco_residual_total = 0
        riscos_modalidade = []
        
        for risco in riscos_comparacao:
            if modalidade in risco['modalidades']:
                fator_mitigacao = risco['modalidades'][modalidade]
                risco_residual = risco['risco_inerente'] * fator_mitigacao
                risco_residual_total += risco_residual
                classificacao_residual, _ = classificar_risco(risco_residual)
                
                dados_comparacao.append({
                    'Modalidade': modalidade,
                    'Risco': risco['risco_chave'],
                    'Risco_Inerente': risco['risco_inerente'],
                    'Fator_Mitigacao': fator_mitigacao,
                    'Risco_Residual': risco_residual,
                    'Classificacao_Residual': classificacao_residual,
                    'Reducao_Percentual': (1 - fator_mitigacao) * 100
                })
                
                riscos_modalidade.append({
                    'risco': risco['risco_chave'],
                    'inerente': risco['risco_inerente'],
                    'fator': fator_mitigacao,
                    'residual': risco_residual
                })
        
        # Armazenar dados acumulados
        risco_acumulado_por_modalidade[modalidade] = {
            'risco_residual_total': risco_residual_total,
            'risco_inerente_total': sum(r['inerente'] for r in riscos_modalidade),
            'eficacia_percentual': ((sum(r['inerente'] for r in riscos_modalidade) - risco_residual_total) / sum(r['inerente'] for r in riscos_modalidade) * 100) if riscos_modalidade else 0,
            'classificacao_total': classificar_risco(risco_residual_total)[0],
            'detalhes': riscos_modalidade
        }
    
    # Visualização do Risco Acumulado
    col1, col2 = st.columns(2)
    
    with col1:
        # Gráfico de barras: Risco residual ACUMULADO por modalidade
        modalidades_nomes = list(risco_acumulado_por_modalidade.keys())
        riscos_acumulados = [dados['risco_residual_total'] for dados in risco_acumulado_por_modalidade.values()]
        eficacias = [dados['eficacia_percentual'] for dados in risco_acumulado_por_modalidade.values()]
        
        df_acumulado = pd.DataFrame({
            'Modalidade': modalidades_nomes,
            'Risco_Residual_Total': riscos_acumulados,
            'Eficacia_Percentual': eficacias
        })
        
        fig_acumulado = px.bar(
            df_acumulado,
            x='Modalidade',
            y='Risco_Residual_Total',
            color='Eficacia_Percentual',
            title="Risco Residual ACUMULADO por Modalidade",
            labels={'Risco_Residual_Total': 'Risco Residual Total'},
            color_continuous_scale='RdYlGn'
        )
        fig_acumulado.update_xaxes(tickangle=45)
        st.plotly_chart(fig_acumulado, use_container_width=True)
    
    with col2:
        # Gráfico de eficácia comparativa
        fig_eficacia = px.bar(
            df_acumulado.sort_values('Eficacia_Percentual', ascending=True),
            x='Eficacia_Percentual',
            y='Modalidade',
            orientation='h',
            title="Eficácia de Mitigação por Modalidade",
            labels={'Eficacia_Percentual': 'Eficácia (%)'},
            color='Eficacia_Percentual',
            color_continuous_scale='RdYlGn'
        )
        st.plotly_chart(fig_eficacia, use_container_width=True)
    
    # Ranking de modalidades baseado no risco acumulado
    st.subheader("🏆 Ranking de Modalidades (Baseado no Risco Acumulado)")
    
    # Preparar dados para ranking
    ranking_data = []
    for modalidade, dados in risco_acumulado_por_modalidade.items():
        ranking_data.append({
            'Modalidade': modalidade,
            'Risco_Residual_Total': dados['risco_residual_total'],
            'Risco_Inerente_Total': dados['risco_inerente_total'],
            'Eficacia_Percentual': dados['eficacia_percentual'],
            'Classificacao_Total': dados['classificacao_total'],
            'Score': dados['eficacia_percentual'] - (dados['risco_residual_total'] / 10)  # Score considerando eficácia e risco residual
        })
    
    df_ranking = pd.DataFrame(ranking_data)
    df_ranking = df_ranking.sort_values('Score', ascending=False)
    df_ranking.index = range(1, len(df_ranking) + 1)
    df_ranking.index.name = 'Posição'
    
    # Colorir o ranking
    def colorir_ranking(row):
        styles = [''] * len(row)
        if 'Classificacao_Total' in row.index:
            pos = row.index.get_loc('Classificacao_Total')
            if row['Classificacao_Total'] == 'Alto':
                styles[pos] = 'background-color: #f8d7da'
            elif row['Classificacao_Total'] == 'Médio':
                styles[pos] = 'background-color: #fff3cd'
            elif row['Classificacao_Total'] == 'Baixo':
                styles[pos] = 'background-color: #d4edda'
        return styles
    
    try:
        st.dataframe(
            df_ranking.style.apply(colorir_ranking, axis=1),
            use_container_width=True
        )
    except:
        st.dataframe(df_ranking, use_container_width=True)
    
    # Insights automáticos
    st.subheader("💡 Insights Automáticos")
    
    if risco_acumulado_por_modalidade:
        melhor_modalidade = min(risco_acumulado_por_modalidade.keys(), 
                               key=lambda x: risco_acumulado_por_modalidade[x]['risco_residual_total'])
        pior_modalidade = max(risco_acumulado_por_modalidade.keys(), 
                             key=lambda x: risco_acumulado_por_modalidade[x]['risco_residual_total'])
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.success(f"""
            **🏆 Melhor Modalidade: {melhor_modalidade}**
            - Risco Residual Total: {risco_acumulado_por_modalidade[melhor_modalidade]['risco_residual_total']:.1f}
            - Eficácia: {risco_acumulado_por_modalidade[melhor_modalidade]['eficacia_percentual']:.1f}%
            - Classificação: {risco_acumulado_por_modalidade[melhor_modalidade]['classificacao_total']}
            """)
        
        with col2:
            st.error(f"""
            **⚠️ Modalidade de Maior Risco: {pior_modalidade}**
            - Risco Residual Total: {risco_acumulado_por_modalidade[pior_modalidade]['risco_residual_total']:.1f}
            - Eficácia: {risco_acumulado_por_modalidade[pior_modalidade]['eficacia_percentual']:.1f}%
            - Classificação: {risco_acumulado_por_modalidade[pior_modalidade]['classificacao_total']}
            """)
        
        # Diferença entre melhor e pior
        diferenca = (risco_acumulado_por_modalidade[pior_modalidade]['risco_residual_total'] - 
                     risco_acumulado_por_modalidade[melhor_modalidade]['risco_residual_total'])
        
        st.info(f"""
        **📊 Análise Comparativa:**
        - Diferença de risco entre melhor e pior modalidade: **{diferenca:.1f} pontos**
        - Escolher a melhor modalidade reduz o risco total em **{diferenca/risco_inerente_total*100:.1f}%**
        - Total de riscos analisados: **{len(riscos_comparacao)}**
        - Risco inerente total (sem mitigação): **{risco_inerente_total:.1f}**
        """)
    
    # Gráfico de composição detalhada
    if not pd.DataFrame(dados_comparacao).empty:
        st.subheader("📈 Mapas de Calor Avançados")
        
        # Criar abas para diferentes visualizações
        tab_heatmap1, tab_heatmap2, tab_composicao = st.tabs([
            "🌡️ Risco Residual", 
            "🎯 Eficácia de Mitigação", 
            "📊 Composição Detalhada"
        ])
        
        with tab_heatmap1:
            # Heatmap de risco residual melhorado
            fig_heatmap_residual = criar_heatmap_modalidades_melhorado(riscos_comparacao)
            st.plotly_chart(fig_heatmap_residual, use_container_width=True)
            st.info("💡 **Interpretação:** Valores menores (verde) indicam menor risco residual. Valores maiores (vermelho) indicam maior risco residual.")
        
        with tab_heatmap2:
            # Heatmap de eficácia melhorado
            fig_heatmap_eficacia = criar_heatmap_eficacia_melhorado(riscos_comparacao)
            st.plotly_chart(fig_heatmap_eficacia, use_container_width=True)
            st.info("💡 **Interpretação:** Valores maiores (verde) indicam maior eficácia na mitigação do risco. Valores menores (vermelho) indicam menor eficácia.")
        
        with tab_composicao:
            # Tabela detalhada de composição
            df_composicao = pd.DataFrame(dados_comparacao)
            df_composicao_pivot = df_composicao.pivot_table(
                index='Risco',
                columns='Modalidade',
                values='Risco_Residual',
                aggfunc='first'
            ).round(1)
            
            st.write("**Tabela de Risco Residual por Modalidade e Risco:**")
            st.dataframe(df_composicao_pivot, use_container_width=True)
            
            # Adicionar linha de totais
            totais_por_modalidade = df_composicao.groupby('Modalidade')['Risco_Residual'].sum().round(1)
            st.write("**Totais por Modalidade:**")
            st.dataframe(totais_por_modalidade.to_frame().T, use_container_width=True)

def dashboard_geral():
    st.header("📈 Dashboard Geral")
    
    if not st.session_state.riscos:
        st.warning("⚠️ Nenhum risco cadastrado.")
        return
    
    # Métricas gerais
    total_riscos = len(st.session_state.riscos)
    riscos_altos = sum(1 for r in st.session_state.riscos if r['classificacao'] == 'Alto')
    riscos_medios = sum(1 for r in st.session_state.riscos if r['classificacao'] == 'Médio')
    riscos_baixos = sum(1 for r in st.session_state.riscos if r['classificacao'] == 'Baixo')
    
    risco_medio_inerente = np.mean([r['risco_inerente'] for r in st.session_state.riscos])
    risco_inerente_total = sum(r['risco_inerente'] for r in st.session_state.riscos)
    
    col1, col2, col3, col4 = st.columns(4)
    
    with col1:
        st.metric("Total de Riscos", total_riscos)
    with col2:
        st.metric("Riscos Altos", riscos_altos, delta=f"{(riscos_altos/total_riscos)*100:.1f}%")
    with col3:
        st.metric("Riscos Médios", riscos_medios, delta=f"{(riscos_medios/total_riscos)*100:.1f}%")
    with col4:
        st.metric("Risco Inerente Total", f"{risco_inerente_total:.1f}")
    
    st.divider()
    
    # Nova seção: Análise de Risco Residual por Modalidade
    st.subheader("🛡️ Análise de Risco Residual Acumulado por Modalidade")
    st.info("💡 **Risco Residual Acumulado** = Soma de todos os riscos residuais para cada modalidade considerando TODOS os riscos.")
    
    # Calcular risco residual acumulado para todas as modalidades
    risco_residual_por_modalidade = {}
    
    for modalidade in st.session_state.modalidades:
        risco_residual_total = 0
        count_riscos = 0
        
        for risco in st.session_state.riscos:
            if modalidade in risco['modalidades']:
                fator_mitigacao = risco['modalidades'][modalidade]
                risco_residual = risco['risco_inerente'] * fator_mitigacao
                risco_residual_total += risco_residual
                count_riscos += 1
        
        if count_riscos > 0:
            eficacia_total = ((sum(r['risco_inerente'] for r in st.session_state.riscos if modalidade in r['modalidades']) - risco_residual_total) / 
                              sum(r['risco_inerente'] for r in st.session_state.riscos if modalidade in r['modalidades']) * 100)
            
            risco_residual_por_modalidade[modalidade] = {
                'risco_residual_total': risco_residual_total,
                'eficacia_percentual': eficacia_total,
                'classificacao': classificar_risco(risco_residual_total)[0],
                'count_riscos': count_riscos
            }
    
    # Visualizar riscos residuais acumulados
    col1, col2 = st.columns(2)
    
    with col1:
        # Métricas de risco residual
        if risco_residual_por_modalidade:
            modalidades_ordenadas = sorted(risco_residual_por_modalidade.items(), 
                                           key=lambda x: x[1]['risco_residual_total'])
            melhor_modalidade_dados = modalidades_ordenadas[0][1]
            melhor_modalidade_nome = modalidades_ordenadas[0][0]
            
            st.success(f"""
            **🏆 Melhor Modalidade (Menor Risco Residual):**
            **{melhor_modalidade_nome}**
            - Risco Residual Total: {melhor_modalidade_dados['risco_residual_total']:.1f}
            - Eficácia: {melhor_modalidade_dados['eficacia_percentual']:.1f}%
            """)
            
            pior_modalidade_dados = modalidades_ordenadas[-1][1]
            pior_modalidade_nome = modalidades_ordenadas[-1][0]
            
            st.error(f"""
            **⚠️ Modalidade de Maior Risco Residual:**
            **{pior_modalidade_nome}**
            - Risco Residual Total: {pior_modalidade_dados['risco_residual_total']:.1f}
            - Eficácia: {pior_modalidade_dados['eficacia_percentual']:.1f}%
            """)
            
            diferenca_risco = (pior_modalidade_dados['risco_residual_total'] - 
                               melhor_modalidade_dados['risco_residual_total'])
            st.info(f"**Diferença de Risco:** {diferenca_risco:.1f} pontos ({diferenca_risco/risco_inerente_total*100:.1f}% do risco total)")
    
    with col2:
        # Gráfico de barras dos riscos residuais acumulados
        if risco_residual_por_modalidade:
            modalidades_nomes = list(risco_residual_por_modalidade.keys())
            riscos_residuais = [dados['risco_residual_total'] for dados in risco_residual_por_modalidade.values()]
            eficacias = [dados['eficacia_percentual'] for dados in risco_residual_por_modalidade.values()]
            
            df_residual = pd.DataFrame({
                'Modalidade': modalidades_nomes,
                'Risco_Residual_Total': riscos_residuais,
                'Eficacia': eficacias
            })
            
            fig_residual = px.bar(
                df_residual.sort_values('Risco_Residual_Total'),
                x='Risco_Residual_Total',
                y='Modalidade',
                orientation='h',
                title="Risco Residual Total por Modalidade",
                color='Eficacia',
                color_continuous_scale='RdYlGn',
                labels={'Risco_Residual_Total': 'Risco Residual Total'}
            )
            st.plotly_chart(fig_residual, use_container_width=True)
    
    # Tabela resumo de todas as modalidades
    if risco_residual_por_modalidade:
        st.subheader("📊 Resumo Executivo - Todas as Modalidades")
        
        dados_resumo = []
        for modalidade, dados in risco_residual_por_modalidade.items():
            dados_resumo.append({
                'Modalidade': modalidade,
                'Risco_Residual_Total': dados['risco_residual_total'],
                'Eficacia_Percentual': dados['eficacia_percentual'],
                'Classificacao_Total': dados['classificacao'],
                'Riscos_Aplicaveis': dados['count_riscos']
            })
        
        df_resumo = pd.DataFrame(dados_resumo)
        df_resumo = df_resumo.sort_values('Risco_Residual_Total')
        df_resumo.index = range(1, len(df_resumo) + 1)
        df_resumo.index.name = 'Ranking'
        
        # Aplicar cores baseado na classificação
        def colorir_classificacao_resumo(row):
            styles = [''] * len(row)
            if 'Classificacao_Total' in row.index:
                pos = row.index.get_loc('Classificacao_Total')
                if row['Classificacao_Total'] == 'Alto':
                    styles[pos] = 'background-color: #f8d7da'
                elif row['Classificacao_Total'] == 'Médio':
                    styles[pos] = 'background-color: #fff3cd'
                elif row['Classificacao_Total'] == 'Baixo':
                    styles[pos] = 'background-color: #d4edda'
            return styles
        
        try:
            st.dataframe(
                df_resumo.style.apply(colorir_classificacao_resumo, axis=1),
                use_container_width=True
            )
        except:
            st.dataframe(df_resumo, use_container_width=True)
    
    st.divider()
    
    # Seções originais
    col1, col2 = st.columns(2)
    
    with col1:
        # Top 5 riscos mais críticos
        st.subheader("🔥 Top 5 Riscos Mais Críticos")
        riscos_ordenados = sorted(st.session_state.riscos, key=lambda x: x['risco_inerente'], reverse=True)[:5]
        
        for i, risco in enumerate(riscos_ordenados, 1):
            classificacao, cor = classificar_risco(risco['risco_inerente'])
            st.write(f"{i}. **{risco['risco_chave']}**")
            st.progress(min(risco['risco_inerente']/100, 1.0))
            st.caption(f"Risco Inerente: {risco['risco_inerente']} ({classificacao})")
            st.write("")
    
    with col2:
        # Eficácia média das modalidades (original, baseado em redução percentual)
        st.subheader("📈 Eficácia Média das Modalidades (Individual)")
        
        eficacia_modalidades = {}
        for modalidade in st.session_state.modalidades:
            reducoes = []
            for risco in st.session_state.riscos:
                if modalidade in risco['modalidades']:
                    fator = risco['modalidades'][modalidade]
                    reducao = (1 - fator) * 100
                    reducoes.append(reducao)
            
            if reducoes:
                eficacia_modalidades[modalidade] = np.mean(reducoes)
        
        if eficacia_modalidades:
            df_eficacia = pd.DataFrame(list(eficacia_modalidades.items()), 
                                       columns=['Modalidade', 'Eficácia (%)'])
            df_eficacia = df_eficacia.sort_values('Eficácia (%)', ascending=True)
            
            fig_eficacia = px.bar(
                df_eficacia,
                x='Eficácia (%)',
                y='Modalidade',
                orientation='h',
                title="Eficácia Média Individual",
                color='Eficácia (%)',
                color_continuous_scale='RdYlGn'
            )
            st.plotly_chart(fig_eficacia, use_container_width=True)
    
    # Matriz de calor consolidada
    st.subheader("🌡️ Matriz de Calor - Todos os Riscos")
    
    try:
        # Criar matriz de posições
        matriz_riscos = np.zeros((11, 11))
        posicoes_riscos = []
        
        for risco in st.session_state.riscos:
            x = min(risco['probabilidade_valor'], 10)
            y = min(risco['impacto_valor'], 10)
            matriz_riscos[y, x] += 1
            posicoes_riscos.append((x, y, risco['risco_chave']))
        
        # Criar heatmap
        fig_matriz = go.Figure(data=go.Heatmap(
            z=matriz_riscos[1:, 1:],  # Excluir linha/coluna 0
            x=list(range(1, 11)),
            y=list(range(1, 11)),
            colorscale='Reds',
            showscale=True
        ))
        
        # Adicionar linhas de grade para delimitar zonas de risco
        fig_matriz.add_hline(y=2.5, line_dash="dash", line_color="blue", opacity=0.5)
        fig_matriz.add_hline(y=5.5, line_dash="dash", line_color="orange", opacity=0.5)
        fig_matriz.add_vline(x=2.5, line_dash="dash", line_color="blue", opacity=0.5)
        fig_matriz.add_vline(x=5.5, line_dash="dash", line_color="orange", opacity=0.5)
        
        # Adicionar anotações para as zonas
        fig_matriz.add_annotation(x=1.5, y=1.5, text="BAIXO", showarrow=False, 
                                 font=dict(size=12, color="green"))
        fig_matriz.add_annotation(x=8, y=8, text="ALTO", showarrow=False, 
                                 font=dict(size=12, color="red"))
        fig_matriz.add_annotation(x=4, y=4, text="MÉDIO", showarrow=False, 
                                 font=dict(size=12, color="orange"))
        
        fig_matriz.update_layout(
            title="Matriz de Calor - Concentração de Riscos",
            xaxis_title="Probabilidade",
            yaxis_title="Impacto",
            width=700,
            height=500
        )
        
        st.plotly_chart(fig_matriz, use_container_width=True)
    except Exception as e:
        st.warning("Erro ao gerar matriz de calor.")
    
    # Insights finais
    if risco_residual_por_modalidade:
        st.subheader("💡 Insights Executivos")
        
        modalidades_ordenadas = sorted(risco_residual_por_modalidade.items(), 
                                       key=lambda x: x[1]['risco_residual_total'])
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.metric(
                "Menor Risco Residual",
                f"{modalidades_ordenadas[0][1]['risco_residual_total']:.1f}",
                delta=f"Modalidade: {modalidades_ordenadas[0][0]}"
            )
        
        with col2:
            st.metric(
                "Maior Risco Residual",
                f"{modalidades_ordenadas[-1][1]['risco_residual_total']:.1f}",
                delta=f"Modalidade: {modalidades_ordenadas[-1][0]}"
            )
        
        with col3:
            amplitude_risco = modalidades_ordenadas[-1][1]['risco_residual_total'] - modalidades_ordenadas[0][1]['risco_residual_total']
            st.metric(
                "Amplitude de Risco",
                f"{amplitude_risco:.1f}",
                delta=f"{amplitude_risco/risco_inerente_total*100:.1f}% do total"
            )

def visualizar_logs():
    st.header("📋 Log de Ações do Sistema")
    
    # Obter logs do banco de dados
    logs = obter_logs()
    
    if not logs:
        st.info("📝 Nenhuma ação registrada ainda.")
        return
    
    # Converter para DataFrame
    df_logs = pd.DataFrame(logs, columns=['Data/Hora', 'Usuário', 'Ação', 'Detalhes'])
    
    # Filtros
    col1, col2 = st.columns(2)
    
    with col1:
        usuarios = df_logs['Usuário'].unique()
        usuario_filtro = st.multiselect(
            "Filtrar por usuário:",
            options=usuarios,
            default=usuarios
        )
    
    with col2:
        acoes = df_logs['Ação'].unique()
        acao_filtro = st.multiselect(
            "Filtrar por ação:",
            options=acoes,
            default=acoes
        )
    
    # Aplicar filtros
    df_filtrado = df_logs[
        (df_logs['Usuário'].isin(usuario_filtro)) & 
        (df_logs['Ação'].isin(acao_filtro))
    ]
    
    # Exibir tabela
    st.dataframe(df_filtrado, use_container_width=True)
    
    # Estatísticas
    st.subheader("📊 Estatísticas de Atividade")
    col1, col2, col3 = st.columns(3)
    
    with col1:
        st.metric("Total de Ações", len(df_filtrado))
    
    with col2:
        acoes_por_usuario = df_filtrado['Usuário'].value_counts()
        st.metric("Ações por Usuário", f"{len(acoes_por_usuario)} usuários")
    
    with col3:
        st.metric("Período Registrado", 
                  f"{df_filtrado['Data/Hora'].min().split()[0]} a {df_filtrado['Data/Hora'].max().split()[0]}")
    
    # Gráfico de atividades por usuário
    fig = px.bar(acoes_por_usuario, 
                 x=acoes_por_usuario.index, 
                 y=acoes_por_usuario.values,
                 title="Ações por Usuário",
                 labels={'x': 'Usuário', 'y': 'Número de Ações'})
    st.plotly_chart(fig, use_container_width=True)

def main():
    # Inicializar banco de dados
    init_db()
    
    # Verificar se o usuário está logado
    if 'user' not in st.session_state:
        st.session_state.user = None
    
    # Se não está logado, mostrar tela de login
    if not st.session_state.user:
        st.title("🔐 Login - Sistema de Gestão de Riscos")
        
        col1, col2, col3 = st.columns([1, 2, 1])
        
        with col2:
            with st.form("login_form"):
                # MUDANÇA: Usar text_input em vez de selectbox para permitir digitação livre
                username = st.text_input("Usuário", placeholder="Digite seu usuário")
                password = st.text_input("Senha", type="password", placeholder="Digite sua senha")
                
                # NOVO: Campo para nome do projeto
                nome_projeto = st.text_input("Nome do Projeto", placeholder="Digite o nome do projeto trabalhado")
                
                submitted = st.form_submit_button("Entrar")
                
                if submitted:
                    if not nome_projeto.strip():
                        st.error("Por favor, digite o nome do projeto")
                    elif verificar_login(username, password):
                        st.session_state.user = username
                        st.session_state.nome_projeto = nome_projeto.strip()
                        st.rerun()
                    else:
                        st.error("Usuário ou senha incorretos")
        
        st.stop()
    
    # Se está logado, mostrar a aplicação normal
    nome_projeto_titulo = st.session_state.get('nome_projeto', 'Projeto')
    st.title(f"🛡️ Dashboard de Avaliação de Riscos - {nome_projeto_titulo}")
    st.markdown(f"*Usuário: {st.session_state.user}*")
    st.markdown("*Metodologia baseada no Roteiro de Auditoria de Gestão de Riscos *")
    
    inicializar_dados()
    
    # Mostrar informações sobre os dados pré-carregados
    if st.session_state.riscos:
        st.success(f"✅ **{len(st.session_state.riscos)} riscos** da planilha foram carregados automaticamente!")
        with st.expander("📋 Visualizar riscos carregados"):
            for i, risco in enumerate(st.session_state.riscos, 1):
                st.write(f"**{i}. {risco['risco_chave']}**")
                st.write(f"    - Risco Inerente: {risco['risco_inerente']} ({risco['classificacao']})")
                st.write(f"    - Impacto: {risco['impacto_valor']} | Probabilidade: {risco['probabilidade_valor']}")
    
    # Sidebar para configurações
    with st.sidebar:
        st.header("⚙️ Configurações")
        
        st.info("💡 **Dados Pré-carregados**\n\nOs riscos da sua planilha já estão carregados! Você pode adicionar novos, editar existentes ou modificar modalidades.")
        
        # Mostrar estatísticas dos riscos
        if st.session_state.riscos:
            st.subheader("📊 Estatísticas Atuais")
            total = len(st.session_state.riscos)
            altos = sum(1 for r in st.session_state.riscos if r['classificacao'] == 'Alto')
            medios = sum(1 for r in st.session_state.riscos if r['classificacao'] == 'Médio')
            baixos = sum(1 for r in st.session_state.riscos if r['classificacao'] == 'Baixo')
            editados = sum(1 for r in st.session_state.riscos if r.get('editado', False))
            adicionados = sum(1 for r in st.session_state.riscos if r.get('personalizado', False))
            
            st.write(f"**Total:** {total} riscos")
            st.write(f"🔴 **Altos:** {altos} ({altos/total*100:.0f}%)")
            st.write(f"🟡 **Médios:** {medios} ({medios/total*100:.0f}%)")
            st.write(f"🟢 **Baixos:** {baixos} ({baixos/total*100:.0f}%)")
            
            if editados > 0:
                st.write(f"✏️ **Personalizados:** {editados}")
            if adicionados > 0:
                st.write(f"➕ **Adicionados:** {adicionados}")
        
        st.divider()
        
        # Gerenciar modalidades
        st.subheader("Modalidades de Mitigação")
        nova_modalidade = st.text_input("Adicionar nova modalidade:")
        if st.button("➕ Adicionar") and nova_modalidade:
            if nova_modalidade not in st.session_state.modalidades:
                st.session_state.modalidades.append(nova_modalidade)
                # Adicionar a nova modalidade a todos os riscos existentes
                for risco in st.session_state.riscos:
                    if 'modalidades' not in risco:
                        risco['modalidades'] = {}
                    risco['modalidades'][nova_modalidade] = 0.5  # Valor padrão
                    # Garante que a nova justificativa também seja adicionada
                    risco.setdefault('justificativas_modalidades', {})[nova_modalidade] = ""
                st.success(f"Modalidade '{nova_modalidade}' adicionada!")
                st.rerun()
            else:
                st.warning("Modalidade já existe!")
        
        # Remover modalidade
        if st.session_state.modalidades:
            modalidade_remover = st.selectbox(
                "Remover modalidade:",
                ["Selecione..."] + st.session_state.modalidades
            )
            if st.button("🗑️ Remover") and modalidade_remover != "Selecione...":
                st.session_state.modalidades.remove(modalidade_remover)
                # Remover a modalidade de todos os riscos
                for risco in st.session_state.riscos:
                    if 'modalidades' in risco and modalidade_remover in risco['modalidades']:
                        del risco['modalidades'][modalidade_remover]
                    if 'justificativas_modalidades' in risco and modalidade_remover in risco['justificativas_modalidades']:
                        del risco['justificativas_modalidades'][modalidade_remover]
                st.success(f"Modalidade '{modalidade_remover}' removida!")
                st.rerun()
        
        st.divider()
        
        # Exportar/Importar dados
        st.subheader("📄 Gerenciar Dados")
        
        # Botão para gerar relatório Word
        if Document and st.button("📄 Gerar Relatório Word", help="Gera relatório completo em formato .docx"):
            with st.spinner("Gerando relatório..."):
                buffer = gerar_relatorio_word()
                if buffer:
                    nome_projeto_arquivo = st.session_state.get('nome_projeto', 'Projeto').replace(' ', '_')
                    st.download_button(
                        label="📥 Baixar Relatório Word",
                        data=buffer,
                        file_name=f"relatorio_riscos_{nome_projeto_arquivo}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="download_report_sidebar"
                    )
                    st.success("✅ Relatório gerado com sucesso!")
        
        if st.button("💾 Exportar dados (JSON)"):
            import json
            dados_export = {
                'riscos': st.session_state.riscos,
                'modalidades': st.session_state.modalidades
            }
            json_string = json.dumps(dados_export, indent=2, ensure_ascii=False)
            st.download_button(
                label="📥 Baixar arquivo JSON",
                data=json_string,
                file_name="avaliacao_riscos.json",
                mime="application/json"
            )
        
        # Resetar dados
        if st.button("🔄 Recarregar dados originais"):
            st.session_state.riscos = []
            st.session_state.modalidades = []
            inicializar_dados()
            st.success("Dados originais recarregados!")
            st.rerun()
        
        if st.button(" Limpar todos os dados"):
            if st.checkbox("⚠️ Confirmo que quero limpar todos os dados"):
                st.session_state.riscos = []
                st.session_state.modalidades = MODALIDADES_PADRAO.copy()
                st.success("Dados limpos!")
                st.rerun()
            else:
                st.warning("Marque a confirmação para limpar os dados")
        
        st.divider()
        st.write(f"Usuário: **{st.session_state.user}**")
        if st.button("🚪 Sair"):
            st.session_state.user = None
            st.rerun()
    
    # Abas principais
    tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs([
        "✏️ Editar Riscos",
        "📝 Cadastro de Riscos",
        "📊 Análise de Riscos", 
        "🔄 Comparação de Modalidades",
        "📈 Dashboard Geral",
        "📋 Log de Ações"
    ])
    
    with tab1:
        editar_riscos()
    
    with tab2:
        cadastro_riscos()
    
    with tab3:
        analise_riscos()
    
    with tab4:
        comparacao_modalidades()
    
    with tab5:
        dashboard_geral()
    
    with tab6:
        visualizar_logs()

if __name__ == "__main__":
    main()
